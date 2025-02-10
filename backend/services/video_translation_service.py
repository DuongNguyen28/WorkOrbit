import moviepy as mp
import speech_recognition as sr
from googletrans import Translator
from docx import Document
import tempfile
import os
from ..services.language_detection_service import LanguageDetectionService

class VideoTranslatorService:
    def __init__(self):
        self.translator = Translator()
        self.recognizer = sr.Recognizer()
        self.language_service = LanguageDetectionService()

    async def translate_text(self, text: str, src_language: str, dest_language: str) -> str:
        """Translate the given text to the specified destination language."""
        translation = await self.translator.translate(text, src=src_language, dest=dest_language)
        return translation.text

    def extract_audio_from_video(self, video_path: str, audio_path: str):
        """Extract audio from the video file."""
        video = mp.VideoFileClip(video_path)
        audio = video.audio
        audio.write_audiofile(audio_path)

    def transcribe_audio_to_text(self, audio_path: str):
        """Transcribe audio to text."""
        with sr.AudioFile(audio_path) as source:
            audio = self.recognizer.record(source)
            try:
                text = self.recognizer.recognize_google(audio)
                return text
            except sr.UnknownValueError:
                return "Could not understand the audio."
            except sr.RequestError:
                return "Error with the speech recognition service."

    def write_translation_to_doc(self, original_text: str, translated_text: str, doc_path: str):
        """Write the original and translated text into a Word document."""
        doc = Document()
        
        doc.add_heading("Original Text", level=1)
        doc.add_paragraph(original_text)
        
        doc.add_paragraph("\n" + "-" * 50 + "\n")
        
        doc.add_heading("Translated Text", level=1)
        doc.add_paragraph(translated_text)
        
        doc.save(doc_path)

    async def process_video_translation(self, file, src_language: str, dest_language: str):
        """Process video translation: extract audio, transcribe, detect language, and translate."""
        warnings = []
        
        # Create temporary files
        with tempfile.NamedTemporaryFile(delete=False, suffix=".mp4") as temp_video:
            video_path = temp_video.name
            temp_video.write(await file.read())

        audio_path = video_path.replace(".mp4", ".wav")
        doc_path = video_path.replace(".mp4", ".docx")
        
        # Process video
        self.extract_audio_from_video(video_path, audio_path)

        text = self.transcribe_audio_to_text(audio_path)

        # Language detection
        detected_language = await self.language_service.detect_language(text)
        if detected_language != src_language:
            warnings.append(f"Warning: Detected language is {detected_language}, but the selected language is {src_language}.")
            # Cleanup temporary files
            os.unlink(video_path)
            os.unlink(audio_path)
            return warnings, None
        
        translated_text = await self.translate_text(text, src_language, dest_language)
        
        self.write_translation_to_doc(text, translated_text, doc_path)
        
        # Cleanup temporary files
        os.unlink(video_path)
        os.unlink(audio_path)
        
        return warnings, doc_path
