import pymupdf
from docx import Document
from docx.shared import Pt
from .translate_client import TranslateClient
from ..services.language_detection_service import LanguageDetectionService
from ..services.elasticsearch_service import ElasticSearchService
import time
import os

class PdfToDocxTranslatorService:
    def __init__(self):
        self.translate_client = TranslateClient()
        self.language_detector = LanguageDetectionService()

    async def translate_text(self, target: str, text: str):
        if isinstance(text, bytes):
            text = text.decode("utf-8")

        result = self.translate_client.translate(text, target_language=target)

        return result["translatedText"]

    async def translate_pdf(self, input_path: str, output_path: str, src_language: str, dest_language: str):
        """Extract text from a PDF, detect its language, compare with the selected source language, and translate."""
        doc = pymupdf.open(input_path)
        document = Document()
        alltext = ""
        warnings = []

        for page_num in range(len(doc)):
            page = doc.load_page(page_num)
            blocks = page.get_text("dict")["blocks"]

            for block in blocks:
                if block['type'] == 0:  # Only process text blocks (not images)
                    for line in block["lines"]:
                        for span in line["spans"]:
                            alltext += span["text"]

                    # Detect the language of the extracted text
                    # detected_language = await self.language_detector.detect_language(alltext)

                    # # Compare detected language with the user's selected source language
                    # if detected_language != src_language:
                    #     warnings.append(f"Warning: Detected language is {detected_language}, but the selected language is {src_language}.")
                    #     return warnings  # Return the warning and stop further processing
                    
                    # Translate the text to the destination language
                    translation = await self.translate_client.translate_text(alltext, dest_language)
                    alltext = ""

                    # Handle text formatting and add it to the Word document
                    font_name = span["font"]
                    font_size = span["size"]

                    p = document.add_paragraph()
                    run = p.add_run(translation)
                    run.font.name = font_name
                    run.font.size = Pt(font_size)

        document.save(output_path)
        
        # If no warnings, return an empty list (indicating no issues)
        return warnings

    async def process_file(self, input_path: str, output_path: str, src_language: str, dest_language: str):
        """Determine file type, detect language, and process accordingly."""
        output_path = f"{src_language}-{dest_language}_{int(time.time())}.docx"
        output_path = os.path.join(os.getcwd(), "backend/misc", output_path)
        if input_path.endswith('.pdf'):
            warnings = await self.translate_pdf(input_path, output_path, src_language, dest_language)
            
            if warnings:
                return {"error": warnings}

            es = ElasticSearchService()
            es.ingest_document(input_path, "pdf")
            es.ingest_document(output_path, "translated_docx")
            
            return {"message": "Translation successful", "file_link": output_path}
        else:
            raise ValueError("Unsupported file format.")
