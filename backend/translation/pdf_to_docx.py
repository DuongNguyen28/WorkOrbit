import asyncio
from googletrans import Translator
import pymupdf
from docx import Document
from docx.shared import Pt

async def translate_text(text: str) -> str:
    translator = Translator()
    # translation = await translator.translate(text, src='en', dest='vi')
    translation = await translator.translate(text, src='ja', dest='vi')

    return translation.text

async def translate_pdf(input_path: str, output_path: str):
    doc = pymupdf.open(input_path)
    document = Document()
    alltext = ""

    # Iterate over all pages in the document
    for page_num in range(len(doc)):
        page = doc.load_page(page_num)
        
        # Extract text as blocks with position info
        blocks = page.get_text("dict")["blocks"]

        # Iterate over each block of text
        for block in blocks:
            if block['type'] == 0:  # Only process text blocks (not images)
                for line in block["lines"]:
                    for span in line["spans"]:
                        alltext += span["text"]
                        
                # # Translate the text
                vietnamese_text = await translate_text(alltext)
                print(vietnamese_text)
                print("\n")

                alltext = ""
                        
                # Handle text formatting: e.g., font, size, and position
                font_name = span["font"]  # The font used
                font_size = span["size"]  # The font size

                # Add text to the Word document with same font and size
                p = document.add_paragraph()
                run = p.add_run(vietnamese_text)
                run.font.name = font_name  # Set font style
                run.font.size = Pt(font_size)  # Set font size
    
    # Save the translated document as a Word file
    document.save(output_path)

async def process_file(input_path: str, output_path: str):
    if input_path.endswith('.pdf'):
        await translate_pdf(input_path, output_path)
    else:
        print("Unsupported file format.")

async def main():
    input_file = "Japanese_test.pdf"
    output_file = "Japanese_test to Vietnamese.docx"
    await process_file(input_file, output_file)

if _name_ == "_main_":
    asyncio.run(main())