from fastapi import FastAPI, UploadFile, File
from fastapi.responses import FileResponse
from fastapi.middleware.cors import CORSMiddleware
import moviepy.editor as mp
import speech_recognition as sr
from googletrans import Translator
from docx import Document
import tempfile
import os

app = FastAPI()

# Enable CORS
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

class VideoTranslator:
    def __init__(self):
        self.translator = Translator()
        self.recognizer = sr.Recognizer()
    
    async def translate_to_vietnamese(self, text: str) -> str:
        translator = Translator()
        translation = await translator.translate(text, src='en', dest='vi')
        return translation.text

    def extract_audio_from_video(self, video_path: str, audio_path: str):
        video = mp.VideoFileClip(video_path)
        audio = video.audio
        audio.write_audiofile(audio_path)

    def transcribe_audio_to_text(self, audio_path: str):
        with sr.AudioFile(audio_path) as source:
            print("Listening for speech...")
            audio = self.recognizer.record(source)
            try:
                print("Recognizing speech...")
                text = self.recognizer.recognize_google(audio)
                return text
            except sr.UnknownValueError:
                return "Could not understand the audio."
            except sr.RequestError:
                return "Error with the speech recognition service."

    def write_translation_to_doc(self, original_text: str, translated_text: str, doc_path: str):
        doc = Document()
        
        # Add original text section
        doc.add_heading("Original Text (English)", level=1)
        doc.add_paragraph(original_text)
        
        # Add separation line
        doc.add_paragraph("\n" + "-" * 50 + "\n")
        
        # Add translated text section
        doc.add_heading("Translated Text (Vietnamese)", level=1)
        doc.add_paragraph(translated_text)
        
        # Save the document
        doc.save(doc_path)

video_translator = VideoTranslator()

@app.post("/translate-video")
async def translate_video(file: UploadFile = File(...)):
    # Create temporary files
    with tempfile.NamedTemporaryFile(delete=False, suffix=".mp4") as temp_video:
        video_path = temp_video.name
        temp_video.write(await file.read())

    audio_path = video_path.replace(".mp4", ".wav")
    doc_path = video_path.replace(".mp4", ".docx")

    # Process video
    video_translator.extract_audio_from_video(video_path, audio_path)
    text = video_translator.transcribe_audio_to_text(audio_path)
    translated_text = await video_translator.translate_to_vietnamese(text)
    video_translator.write_translation_to_doc(text, translated_text, doc_path)

    # Cleanup temporary files
    os.unlink(video_path)
    os.unlink(audio_path)

    return FileResponse(doc_path, media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document", filename="translated_document.docx")
